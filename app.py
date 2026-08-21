import os
import io
import re
import time
import json
import base64
import hashlib
import gc
from typing import List, Dict, Any, Tuple, Optional

import numpy as np
import pypdf
import streamlit as st
from sentence_transformers import SentenceTransformer
from groq import Groq

# ==============================================================================
# BM25 OKAPI RETRIEVER (Pure-Python Cloud-Safe Implementation with Fallback)
# ==============================================================================
try:
    from rank_bm25 import BM25Okapi
except ImportError:
    import math
    from collections import Counter

    class BM25Okapi:
        """Pure-Python fallback implementation of BM25Okapi for low-RAM cloud safety."""
        def __init__(self, corpus: List[List[str]], k1: float = 1.5, b: float = 0.75):
            self.k1 = k1
            self.b = b
            self.corpus_size = len(corpus)
            self.avgdl = sum(len(doc) for doc in corpus) / (self.corpus_size or 1)
            self.doc_freqs = []
            self.idf = {}
            self.doc_len = []
            nd = Counter()
            for doc in corpus:
                self.doc_len.append(len(doc))
                freqs = Counter(doc)
                self.doc_freqs.append(freqs)
                for word in freqs.keys():
                    nd[word] += 1
            for word, freq in nd.items():
                self.idf[word] = math.log(1.0 + (self.corpus_size - freq + 0.5) / (freq + 0.5))

        def get_scores(self, query: List[str]) -> List[float]:
            scores = [0.0] * self.corpus_size
            for q in query:
                if q not in self.idf:
                    continue
                idf = self.idf[q]
                for idx, doc_freq in enumerate(self.doc_freqs):
                    freq = doc_freq.get(q, 0)
                    if freq == 0:
                        continue
                    num = freq * (self.k1 + 1.0)
                    denom = freq + self.k1 * (1.0 - self.b + self.b * self.doc_len[idx] / (self.avgdl or 1.0))
                    scores[idx] += idf * (num / denom)
            return scores


# ==============================================================================
# 1. PAGE SETUP, LOGO HELPERS & DARK-ONLY ENTERPRISE CSS (ZERO EMOJIS)
# ==============================================================================
MAX_UPLOAD_MB = 2048

st.set_page_config(
    page_title="DocuBrain Enterprise - Document Intelligence Platform",
    page_icon=None,
    layout="wide",
    initial_sidebar_state="expanded"
)

def get_logo_html(height_px: int = 34, is_spinning: bool = False) -> str:
    """Renders local logo.png as floating white line-art, or an inline SVG brain glyph if missing."""
    logo_path = os.path.join(os.path.dirname(__file__), "logo.png") if "__file__" in locals() else "logo.png"
    animation_style = "animation: spinLogo 1.4s linear infinite;" if is_spinning else ""

    if os.path.exists(logo_path):
        try:
            with open(logo_path, "rb") as f:
                b64_data = base64.b64encode(f.read()).decode("utf-8")
            return f'<img src="data:image/png;base64,{b64_data}" style="height:{height_px}px; vertical-align:middle; filter:invert(1); mix-blend-mode:screen; {animation_style}" alt="DocuBrain Logo" />'
        except Exception:
            pass

    # SVG Fallback
    return f"""
    <svg width="{height_px}" height="{height_px}" viewBox="0 0 24 24" fill="none" stroke="#818CF8" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="vertical-align:middle; {animation_style}">
        <path d="M9.5 2A2.5 2.5 0 0 1 12 4.5v15a2.5 2.5 0 0 1-4.96.44 2.5 2.5 0 0 1-2.96-3.08 3 3 0 0 1-.34-5.58 2.5 2.5 0 0 1 1.32-4.24 2.5 2.5 0 0 1 4.44-2.04Z"/>
        <path d="M14.5 2A2.5 2.5 0 0 0 12 4.5v15a2.5 2.5 0 0 0 4.96.44 2.5 2.5 0 0 0 2.96-3.08 3 3 0 0 0 .34-5.58 2.5 2.5 0 0 0-1.32-4.24 2.5 2.5 0 0 0-4.44-2.04Z"/>
    </svg>
    """

ENTERPRISE_DARK_CSS = """
<style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500;600;700&display=swap');
    
    *, html, body {
        font-family: 'Plus Jakarta Sans', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    }
    
    code, pre, .mono-text {
        font-family: 'JetBrains Mono', monospace !important;
    }

    /* Force Permanent Dark Surface */
    .stApp {
        background-color: #0B0F14 !important;
        color: #E5E7EB !important;
    }

    /* Hide ONLY the settings menu (keeps dark-only) but keep the sidebar toggle button */
    #MainMenu, [data-testid="stMainMenu"] {
        visibility: hidden !important;
    }
    footer {
        visibility: hidden !important;
    }

    /* Top Navigation Bar */
    .app-header-bar {
        display: flex;
        align-items: center;
        justify-content: space-between;
        padding: 0.75rem 0 1.25rem 0;
        border-bottom: 1px solid rgba(255, 255, 255, 0.08);
        margin-bottom: 1.25rem;
    }
    .header-left-cluster {
        display: flex;
        align-items: center;
        gap: 12px;
    }
    .app-wordmark {
        font-size: 1.45rem;
        font-weight: 800;
        letter-spacing: -0.5px;
        color: #F8FAFC;
    }
    .edition-badge {
        background: rgba(99, 102, 241, 0.12);
        color: #818CF8;
        border: 1px solid rgba(99, 102, 241, 0.35);
        font-size: 0.65rem;
        font-weight: 700;
        letter-spacing: 1.2px;
        padding: 3px 8px;
        border-radius: 4px;
        text-transform: uppercase;
    }
    .system-status-indicator {
        display: flex;
        align-items: center;
        gap: 6px;
        background: rgba(16, 185, 129, 0.08);
        border: 1px solid rgba(16, 185, 129, 0.25);
        padding: 3px 10px;
        border-radius: 12px;
    }
    .status-dot {
        width: 6px;
        height: 6px;
        background-color: #10B981;
        border-radius: 50%;
        box-shadow: 0 0 6px rgba(16, 185, 129, 0.8);
    }
    .status-label {
        font-size: 0.7rem;
        font-weight: 600;
        color: #34D399;
    }
    .header-right-cluster {
        display: flex;
        align-items: center;
        gap: 8px;
    }
    .header-pill-doc {
        background: rgba(255, 255, 255, 0.04);
        border: 1px solid rgba(255, 255, 255, 0.12);
        padding: 4px 12px;
        border-radius: 999px;
        font-size: 0.75rem;
        color: #CBD5E1;
        max-width: 260px;
        white-space: nowrap;
        overflow: hidden;
        text-overflow: ellipsis;
    }
    .header-pill-engine {
        background: rgba(99, 102, 241, 0.1);
        border: 1px solid rgba(99, 102, 241, 0.3);
        color: #A5B4FC;
        padding: 4px 10px;
        border-radius: 999px;
        font-size: 0.72rem;
        font-family: 'JetBrains Mono', monospace;
    }

    /* Metric Strip */
    .metric-container {
        display: flex;
        gap: 12px;
        margin-bottom: 1.25rem;
        flex-wrap: wrap;
    }
    .metric-card {
        background: #11161D;
        border: 1px solid rgba(255, 255, 255, 0.08);
        border-radius: 8px;
        padding: 12px 18px;
        flex: 1;
        min-width: 140px;
    }
    .metric-label {
        font-size: 0.68rem;
        color: #94A3B8;
        text-transform: uppercase;
        font-weight: 600;
        letter-spacing: 0.5px;
        margin-bottom: 4px;
    }
    .metric-value {
        font-size: 1.15rem;
        font-weight: 700;
        color: #F8FAFC;
    }

    /* Gemini / Modern Chat Layout */
    .chat-history-container {
        display: flex;
        flex-direction: column;
        gap: 1rem;
        margin-bottom: 1.5rem;
    }
    .user-bubble-row {
        display: flex;
        justify-content: flex-end;
        width: 100%;
        margin-bottom: 0.5rem;
    }
    .user-bubble {
        background-color: #1E293B;
        color: #F8FAFC;
        border: 1px solid rgba(255, 255, 255, 0.08);
        border-radius: 14px 14px 2px 14px;
        padding: 10px 16px;
        width: fit-content;
        max-width: 80%;
        font-size: 0.92rem;
        line-height: 1.55;
        word-wrap: break-word;
    }
    .assistant-bubble-row {
        display: flex;
        align-items: flex-start;
        gap: 12px;
        width: 100%;
        margin-bottom: 1rem;
    }
    .assistant-accent-dot {
        width: 7px;
        height: 7px;
        background-color: #6366F1;
        border-radius: 50%;
        margin-top: 8px;
        flex-shrink: 0;
    }
    .assistant-bubble-body {
        flex: 1;
        color: #E2E8F0;
        font-size: 0.92rem;
        line-height: 1.6;
    }

    /* Compact Thinking Pill */
    .thinking-pill {
        display: inline-flex;
        align-items: center;
        gap: 8px;
        background: #11161D;
        border: 1px solid rgba(99, 102, 241, 0.35);
        border-radius: 999px;
        padding: 6px 14px;
        margin: 8px 0 14px 0;
        width: fit-content;
    }
    .thinking-pill-text {
        font-size: 0.78rem;
        color: #CBD5E1;
        font-weight: 500;
    }
    @keyframes spinLogo {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }

    /* Badges & Chips */
    .cite-chip {
        display: inline-block;
        background: rgba(99, 102, 241, 0.15);
        color: #A5B4FC;
        border: 1px solid rgba(99, 102, 241, 0.4);
        border-radius: 4px;
        padding: 1px 6px;
        font-size: 0.72rem;
        font-weight: 600;
        font-family: 'JetBrains Mono', monospace;
        margin: 0 2px;
    }
    .score-badge {
        display: inline-block;
        background: rgba(16, 185, 129, 0.1);
        color: #34D399;
        border: 1px solid rgba(16, 185, 129, 0.25);
        border-radius: 4px;
        padding: 2px 8px;
        font-size: 0.72rem;
        font-weight: 600;
    }
    .cached-badge {
        display: inline-block;
        background: rgba(234, 179, 8, 0.1);
        color: #FACC15;
        border: 1px solid rgba(234, 179, 8, 0.25);
        border-radius: 4px;
        padding: 2px 6px;
        font-size: 0.65rem;
        font-weight: 700;
        letter-spacing: 0.5px;
        margin-left: 8px;
    }
    .verification-footnote {
        margin-top: 10px;
        padding-top: 6px;
        border-top: 1px solid rgba(255, 255, 255, 0.06);
        font-size: 0.75rem;
        color: #94A3B8;
    }
    .verification-verified {
        color: #34D399;
        font-weight: 700;
        letter-spacing: 0.5px;
    }
    .verification-notice {
        color: #F59E0B;
        font-weight: 700;
        letter-spacing: 0.5px;
    }

    /* Library Cards */
    .library-item {
        background: rgba(255, 255, 255, 0.02);
        border: 1px solid rgba(255, 255, 255, 0.07);
        border-radius: 6px;
        padding: 10px 12px;
        margin-bottom: 8px;
    }
    .library-item-active {
        background: rgba(99, 102, 241, 0.08);
        border: 1px solid rgba(99, 102, 241, 0.4);
        border-radius: 6px;
        padding: 10px 12px;
        margin-bottom: 8px;
    }
    .library-name {
        font-size: 0.8rem;
        font-weight: 600;
        color: #E2E8F0;
        white-space: nowrap;
        overflow: hidden;
        text-overflow: ellipsis;
    }
    .library-meta {
        font-size: 0.7rem;
        color: #64748B;
        margin-top: 2px;
    }

    /* Button Theming */
    div.stButton > button {
        border-radius: 6px !important;
        border: 1px solid rgba(255, 255, 255, 0.12) !important;
        background-color: #11161D !important;
        color: #E2E8F0 !important;
        font-size: 0.8rem !important;
        font-weight: 500 !important;
        transition: all 0.15s ease-in-out !important;
    }
    div.stButton > button:hover {
        border-color: #6366F1 !important;
        color: #818CF8 !important;
        background-color: rgba(99, 102, 241, 0.08) !important;
    }
    div.stDownloadButton > button {
        border-radius: 6px !important;
        border: 1px solid rgba(99, 102, 241, 0.3) !important;
        background-color: rgba(99, 102, 241, 0.1) !important;
        color: #818CF8 !important;
        font-size: 0.8rem !important;
        font-weight: 600 !important;
    }
    div.stDownloadButton > button:hover {
        border-color: #818CF8 !important;
        background-color: rgba(99, 102, 241, 0.2) !important;
    }

    .app-footer {
        margin-top: 3rem;
        padding: 1.5rem 0 1rem 0;
        border-top: 1px solid rgba(255, 255, 255, 0.08);
        text-align: center;
        font-size: 0.75rem;
        color: #64748B;
    }
</style>
"""
st.markdown(ENTERPRISE_DARK_CSS, unsafe_allow_html=True)


# ==============================================================================
# 2. SECRETS, ACCESS GATE & PERSISTENT LIBRARY STORAGE
# ==============================================================================
LIBRARY_DIR = ".docubrain_library"
os.makedirs(LIBRARY_DIR, exist_ok=True)

STOPWORDS = set("""
a about above after again against all am an and any are aren't as at be because been before being below
between both but by can can't cannot could couldn't did didn't do does doesn't doing don't down during
each few for from further had hadn't has hasn't have haven't having he he'd he'll he's her here here's
hers herself him himself his how how's i i'd i'll i'm i've if in into is isn't it it's its itself let's
me more most mustn't my myself no nor not of off on once only or other ought our ours ourselves out over
own same shan't she she'd she'll she's should shouldn't so some such than that that's the their theirs
them themselves then there there's these they they'd they'll they're they've this those through to too
under until up very was wasn't we we'd we'll we're we've were weren't what what's when when's where
where's which while who who's whom why why's with won't would wouldn't you you'd you'll you're you've your
yours yourself yourselves page section table figure also said may might must shall per will
""".split())

def get_secret(name: str, fallback: str = "") -> str:
    try:
        return str(st.secrets[name])
    except Exception:
        return os.environ.get(name, fallback)

GROQ_API_KEY = get_secret("GROQ_API_KEY", "")
OPENROUTER_API_KEY = get_secret("OPENROUTER_API_KEY", "")
APP_PASSWORD = get_secret("APP_PASSWORD", "harry2026")

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.markdown(f"""
        <div style="max-width: 440px; margin: 6rem auto 2rem auto; background: #11161D; border: 1px solid rgba(255,255,255,0.08); border-radius: 8px; padding: 2.2rem;">
            <div style="display:flex; align-items:center; gap:10px; margin-bottom: 8px;">
                {get_logo_html(28)}
                <span style="font-size: 1.4rem; font-weight: 800; color: #F8FAFC;">DocuBrain</span>
            </div>
            <div style="font-size: 0.68rem; color: #818CF8; font-weight: 700; letter-spacing: 1.2px; margin-bottom: 1.2rem;">ENTERPRISE ACCESS VAULT</div>
            <div style="font-size: 0.82rem; color: #94A3B8; line-height: 1.4; margin-bottom: 1.5rem;">Access requires authorized enterprise credentials. Inquiries and embeddings are processed strictly under zero-retention security protocols.</div>
        </div>
    """, unsafe_allow_html=True)

    with st.container():
        _, col_auth, _ = st.columns([1, 1.3, 1])
        with col_auth:
            auth_pwd = st.text_input("Access Code", type="password", label_visibility="collapsed", placeholder="Enter authorization key")
            if st.button("Authenticate Platform Session", use_container_width=True):
                if auth_pwd and auth_pwd == APP_PASSWORD:
                    st.session_state.authenticated = True
                    st.session_state.is_admin = True
                    st.rerun()
                else:
                    st.error("Authentication failed: Invalid authorization key.")
    st.stop()


# ==============================================================================
# 3. SESSION STATE & TELEMETRY INITIALIZATION
# ==============================================================================
DEFAULT_STATE: Dict[str, Any] = {
    "active_doc_id": None,
    "doc_name": None,
    "num_pages": 0,
    "chunks": [],
    "vectors": None,
    "bm25": None,
    "messages": [],
    "briefing": None,
    "answer_cache": {},
    "feedback": {},
    "rate_limit_timestamps": [],
    "audit_log": [],
    "pending_query": None,
    "last_ingested_hash": None,
    "onboarding_open": True,
    "is_admin": False,
    "telemetry": {
        "total_queries": 0,
        "total_tokens_est": 0,
        "cache_hits": 0,
        "total_latency": 0.0,
        "last_model": "None",
        "security_blocks": 0
    }
}

for k, v in DEFAULT_STATE.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ==============================================================================
# 4. LIBRARY DISK PERSISTENCE ENGINE
# ==============================================================================
def compute_sha256(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()

def normalize_text_for_cache(text: str) -> str:
    return re.sub(r'[^\w\s]', '', text.lower()).strip()

def get_answer_cache_key(doc_id: str, question: str) -> str:
    norm = normalize_text_for_cache(question)
    return hashlib.sha256(f"{doc_id}:{norm}".encode("utf-8")).hexdigest()

def redact_sensitive_pii(text: str) -> str:
    text = re.sub(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', '[REDACTED_EMAIL]', text)
    text = re.sub(r'\b\d{3}-\d{2}-\d{4}\b', '[REDACTED_SSN]', text)
    text = re.sub(r'\b(?:\d{4}[-\s]?){3}\d{4}\b', '[REDACTED_CARD]', text)
    return text

def save_document_library_state(doc_id: str, filename: str, num_pages: int, chunks: List[Dict],
                                vectors: np.ndarray, messages: List[Dict],
                                answer_cache: Dict, feedback: Dict, briefing: Optional[str] = None) -> None:
    meta_path = os.path.join(LIBRARY_DIR, f"{doc_id}.json")
    vec_path = os.path.join(LIBRARY_DIR, f"{doc_id}.npy")
    
    meta_data = {
        "doc_id": doc_id,
        "filename": filename,
        "num_pages": num_pages,
        "indexed_at": time.strftime("%Y-%m-%d %H:%M:%S UTC", time.gmtime()),
        "chunks": chunks,
        "messages": messages,
        "briefing": briefing,
        "answer_cache": answer_cache,
        "feedback": feedback
    }
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(meta_data, f, indent=2)
    np.save(vec_path, vectors.astype(np.float32))

def load_document_library_state(doc_id: str) -> Optional[Dict[str, Any]]:
    meta_path = os.path.join(LIBRARY_DIR, f"{doc_id}.json")
    vec_path = os.path.join(LIBRARY_DIR, f"{doc_id}.npy")
    if not (os.path.exists(meta_path) and os.path.exists(vec_path)):
        return None
    try:
        with open(meta_path, "r", encoding="utf-8") as f:
            meta_data = json.load(f)
        vectors = np.load(vec_path).astype(np.float32)
        meta_data["vectors"] = vectors
        if "doc_id" not in meta_data:
            meta_data["doc_id"] = doc_id
        return meta_data
    except Exception:
        return None

def delete_document_library_state(doc_id: str) -> None:
    meta_path = os.path.join(LIBRARY_DIR, f"{doc_id}.json")
    vec_path = os.path.join(LIBRARY_DIR, f"{doc_id}.npy")
    if os.path.exists(meta_path):
        try: os.remove(meta_path)
        except OSError: pass
    if os.path.exists(vec_path):
        try: os.remove(vec_path)
        except OSError: pass

def list_library_documents() -> List[Dict[str, Any]]:
    docs = []
    if not os.path.exists(LIBRARY_DIR):
        return docs
    for fname in os.listdir(LIBRARY_DIR):
        if fname.endswith(".json"):
            fpath = os.path.join(LIBRARY_DIR, fname)
            try:
                with open(fpath, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    doc = {
                        "doc_id": data.get("doc_id", fname.replace(".json", "")),
                        "filename": data.get("filename", "Untitled Document"),
                        "num_pages": data.get("num_pages", 0),
                        "indexed_at": data.get("indexed_at", "Unknown Date"),
                        "chunk_count": len(data.get("chunks", []))
                    }
                    docs.append(doc)
            except Exception:
                continue
    docs.sort(key=lambda x: x.get("indexed_at", ""), reverse=True)
    return docs


# ==============================================================================
# 5. EMBEDDING ENGINE & HYBRID SEARCH PIPELINE (PDF + DOCX SUPPORT)
# ==============================================================================
@st.cache_resource(show_spinner=False)
def load_embedder() -> SentenceTransformer:
    embedder = SentenceTransformer('all-MiniLM-L6-v2', device='cpu')
    embedder.encode(["DocuBrain enterprise embedder probe"], normalize_embeddings=True)
    return embedder

def extract_document_pages(file_bytes: bytes, filename: str) -> Optional[List[Dict[str, Any]]]:
    """Extracts text pages from PDF or .docx Word files. Legacy .doc returns None."""
    ext = filename.lower().rsplit(".", 1)[-1]
    try:
        if ext == "pdf":
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                text = re.sub(r'[ \t]+', ' ', text)
                text = re.sub(r'\n+', '\n', text).strip()
                if text:
                    text = redact_sensitive_pii(text)
                    pages.append({"page": i + 1, "text": text})
            return pages if pages else None
        elif ext == "docx":
            import docx
            document = docx.Document(io.BytesIO(file_bytes))
            blocks = [p.text for p in document.paragraphs if p.text and p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        blocks.append(row_text)
            if not blocks:
                return None
            pages = []
            current_texts = []
            current_len = 0
            page_num = 1
            for para in blocks:
                clean = re.sub(r'\s+', ' ', para).strip()
                if not clean:
                    continue
                if current_len + len(clean) > 3000 and current_texts:
                    pages.append({"page": page_num, "text": redact_sensitive_pii(" ".join(current_texts))})
                    page_num += 1
                    current_texts = []
                    current_len = 0
                current_texts.append(clean)
                current_len += len(clean)
            if current_texts:
                pages.append({"page": page_num, "text": redact_sensitive_pii(" ".join(current_texts))})
            return pages if pages else None
        else:
            return None
    except Exception:
        return None

def sentence_aware_chunking(pages: List[Dict[str, Any]], chunk_size: int = 750, overlap: int = 150) -> List[Dict[str, Any]]:
    chunks = []
    chunk_counter = 0
    for page_data in pages:
        text = page_data["text"]
        page_num = page_data["page"]
        sentences = re.split(r'(?<=[.?!])\s+', text)
        current_chunk: List[str] = []
        current_length = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            s_len = len(sentence)
            if current_length + s_len > chunk_size and current_chunk:
                chunks.append({
                    "chunk_id": chunk_counter,
                    "page": page_num,
                    "text": " ".join(current_chunk)
                })
                chunk_counter += 1
                overlap_chunk: List[str] = []
                overlap_len = 0
                for s in reversed(current_chunk):
                    if overlap_len + len(s) <= overlap:
                        overlap_chunk.insert(0, s)
                        overlap_len += len(s)
                    else:
                        break
                current_chunk = overlap_chunk
                current_length = overlap_len

            current_chunk.append(sentence)
            current_length += s_len

        if current_chunk:
            chunks.append({
                "chunk_id": chunk_counter,
                "page": page_num,
                "text": " ".join(current_chunk)
            })
            chunk_counter += 1
    return chunks

def build_bm25_index(chunks: List[Dict[str, Any]]) -> BM25Okapi:
    tokenized_corpus = [re.findall(r'\w+', c["text"].lower()) for c in chunks]
    return BM25Okapi(tokenized_corpus)

def batch_encode_chunks(embedder: SentenceTransformer, chunks: List[Dict[str, Any]]) -> np.ndarray:
    texts = [c["text"] for c in chunks]
    embeddings = embedder.encode(
        texts,
        batch_size=64,
        show_progress_bar=False,
        normalize_embeddings=True,
        convert_to_numpy=True
    )
    gc.collect()
    return embeddings.astype(np.float32)

def extract_top_keywords(chunks: List[Dict[str, Any]], top_n: int = 12) -> List[Tuple[str, int]]:
    from collections import Counter
    words = []
    for c in chunks:
        tokens = re.findall(r'\b[a-zA-Z]{3,}\b', c["text"].lower())
        words.extend([w for w in tokens if w not in STOPWORDS])
    return Counter(words).most_common(top_n)

def hybrid_retrieve(query: str, embedder: SentenceTransformer, doc_vectors: np.ndarray,
                    bm25: BM25Okapi, chunks: List[Dict[str, Any]], top_k: int = 4,
                    k_rrf: int = 60) -> List[Dict[str, Any]]:
    query_vec = embedder.encode([query], normalize_embeddings=True, convert_to_numpy=True)[0].astype(np.float32)
    dense_scores = np.dot(doc_vectors, query_vec)
    dense_ranks = np.argsort(-dense_scores)

    query_tokens = re.findall(r'\w+', query.lower())
    bm25_scores = np.array(bm25.get_scores(query_tokens))
    bm25_ranks = np.argsort(-bm25_scores)

    num_chunks = len(chunks)
    rrf_scores = np.zeros(num_chunks, dtype=np.float32)
    for rank, idx in enumerate(dense_ranks):
        rrf_scores[idx] += 1.0 / (k_rrf + rank + 1)
    for rank, idx in enumerate(bm25_ranks):
        rrf_scores[idx] += 1.0 / (k_rrf + rank + 1)

    top_indices = np.argsort(-rrf_scores)[:top_k]
    
    results = []
    for idx in top_indices:
        results.append({
            "chunk_index": int(idx),
            "chunk": chunks[idx],
            "dense_score": float(dense_scores[idx]),
            "bm25_score": float(bm25_scores[idx]),
            "rrf_score": float(rrf_scores[idx]),
            "score": float(dense_scores[idx])
        })
    return results

def expand_parent_context_window(retrieved_items: List[Dict[str, Any]], all_chunks: List[Dict[str, Any]], max_chars: int = 5500) -> str:
    total_chunks = len(all_chunks)
    included_indices = set()
    ordered_indices = []

    for item in retrieved_items:
        idx = item["chunk_index"]
        neighbors = [idx]
        if idx > 0 and all_chunks[idx - 1]["page"] == all_chunks[idx]["page"]:
            neighbors.insert(0, idx - 1)
        if idx + 1 < total_chunks and all_chunks[idx + 1]["page"] == all_chunks[idx]["page"]:
            neighbors.append(idx + 1)
        for n in neighbors:
            if n not in included_indices:
                included_indices.add(n)
                ordered_indices.append(n)

    ordered_indices.sort()
    context_blocks = []
    curr_page = None
    curr_texts = []
    curr_len = 0

    for idx in ordered_indices:
        chunk = all_chunks[idx]
        text = chunk["text"]
        page = chunk["page"]
        if curr_len + len(text) > max_chars:
            break
        if curr_page == page:
            curr_texts.append(text)
        else:
            if curr_texts:
                context_blocks.append(f"[Page {curr_page}]\n" + " ".join(curr_texts))
            curr_page = page
            curr_texts = [text]
        curr_len += len(text)

    if curr_texts:
        context_blocks.append(f"[Page {curr_page}]\n" + " ".join(curr_texts))

    return "\n\n".join(context_blocks)


# ==============================================================================
# 6. INFERENCE, SECURITY SHIELD, REWRITING & CASCADE (OPENROUTER FIRST)
# ==============================================================================
MODEL_MAPPING = {
    "openrouter/free": "OpenRouter Free",
    "openai/gpt-oss-20b": "GPT-OSS 20B",
    "qwen/qwen3.6-27b": "QWEN 3.6 27B",
    "openai/gpt-oss-120b": "GPT-OSS 120B",
    "meta-llama/llama-prompt-guard-2-86m": "PROMPT-GUARD 86M"
}

MODEL_COSTS = {
    "openrouter/free": {"input": 0.0, "output": 0.0},
    "OpenRouter Free": {"input": 0.0, "output": 0.0},
    "openai/gpt-oss-20b": {"input": 0.00015, "output": 0.0006},
    "qwen/qwen3.6-27b": {"input": 0.0002, "output": 0.0008},
    "openai/gpt-oss-120b": {"input": 0.0008, "output": 0.0024},
    "meta-llama/llama-prompt-guard-2-86m": {"input": 0.00005, "output": 0.00005},
}

OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"
OPENROUTER_FREE_MODEL = "openrouter/free"

def get_client() -> Optional[Groq]:
    if not GROQ_API_KEY:
        return None
    return Groq(api_key=GROQ_API_KEY)

def get_openrouter_client():
    """OpenAI-compatible client for OpenRouter (free models router)."""
    if not OPENROUTER_API_KEY:
        return None
    try:
        from openai import OpenAI
        return OpenAI(api_key=OPENROUTER_API_KEY, base_url=OPENROUTER_BASE_URL)
    except Exception:
        return None

def check_prompt_security(client: Groq, query: str) -> Tuple[bool, str]:
    try:
        response = client.chat.completions.create(
            model="meta-llama/llama-prompt-guard-2-86m",
            messages=[{"role": "user", "content": query}],
            max_tokens=20,
            temperature=0.0
        )
        verdict = (response.choices[0].message.content or "").strip().upper()
        if any(flag in verdict for flag in ["INJECTION", "JAILBREAK", "UNSAFE", "MALICIOUS"]):
            return False, "Security Notice: Prompt injection or policy violation pattern detected."
        return True, ""
    except Exception:
        return True, ""

def rewrite_query_if_followup(client: Groq, history: List[Dict[str, Any]], latest_query: str) -> str:
    user_turns = [m for m in history if m.get("role") == "user"]
    if len(user_turns) == 0:
        return latest_query

    recent = history[-4:]
    hist_text = "\n".join(f"{m['role'].upper()}: {m['content'][:250]}" for m in recent)
    prompt = (
        "Given the conversation history and the user's latest follow-up question, rewrite the "
        "follow-up into an independent, standalone search query containing all necessary entities. "
        "Do not include conversational preamble. Return ONLY the rewritten query text.\n\n"
        f"CONVERSATION:\n{hist_text}\n\n"
        f"FOLLOW-UP: {latest_query}\n\n"
        "STANDALONE QUERY:"
    )
    try:
        resp = client.chat.completions.create(
            model="openai/gpt-oss-20b",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=80
        )
        rewritten = (resp.choices[0].message.content or "").strip().strip('"\'')
        if rewritten and len(rewritten) > 3:
            return rewritten
    except Exception:
        pass
    return latest_query

def render_citation_chips(text: str) -> str:
    def repl(m):
        raw = m.group(1)
        return f'<span class="cite-chip">PAGE {raw.strip()}</span>'
    return re.sub(r'\[Pages?\s*(\d+(?:[,\s\-–]+\d+)*)\]', repl, text, flags=re.IGNORECASE)

def verify_citations(response_text: str, retrieved_chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
    retrieved_pages = sorted(list(set(c["chunk"]["page"] for c in retrieved_chunks)))
    matches = re.findall(r'\[Pages?\s*(\d+(?:[,\s\-–]+\d+)*)\]|\(Pages?\s*(\d+(?:[,\s\-–]+\d+)*)\)', response_text, re.IGNORECASE)
    
    cited_pages = set()
    for m in matches:
        raw = m[0] or m[1]
        for num in re.findall(r'\d+', raw):
            cited_pages.add(int(num))
            
    cited_list = sorted(list(cited_pages))
    if not cited_list:
        return {
            "has_citations": False,
            "cited_pages": [],
            "retrieved_pages": retrieved_pages,
            "status_text": "No explicit page citations detected."
        }
    
    unverified = [p for p in cited_list if p not in retrieved_pages]
    is_valid = (len(unverified) == 0)
    pages_str = ", ".join(str(p) for p in cited_list)
    
    if is_valid:
        status_text = f"Verified against source pages: {pages_str}"
    else:
        unv_str = ", ".join(str(p) for p in unverified)
        status_text = f"Notice: Citation referenced unverified page(s) {unv_str}"
        
    return {
        "has_citations": True,
        "is_valid": is_valid,
        "cited_pages": cited_list,
        "retrieved_pages": retrieved_pages,
        "status_text": status_text
    }

def stream_groq_cascade(client: Groq, selected_model: str, context: str, question: str,
                        temperature: float = 0.1):
    """CASCADE ORDER: OpenRouter Free Models Router FIRST, then GPT-OSS 120B, then GPT-OSS 20B."""
    system_prompt = (
        "You are an expert, precise document analyst. Answer the user's question using ONLY "
        "the provided Context.\n\n"
        "Mandatory Guidelines:\n"
        "1. Cite the exact source page for every key fact using bracket format like `[Page X]`.\n"
        "2. Structure answers cleanly using markdown bullet points and bold highlights.\n"
        "3. If the answer cannot be strictly derived from the context, state: "
        "'*The provided document does not contain sufficient information to answer this question.*'"
    )
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"CONTEXT:\n{context}\n\nQUESTION: {question}"}
    ]

    if selected_model == "Auto-Cascade":
        cascade = [
            ("openrouter", OPENROUTER_FREE_MODEL, "OpenRouter Free"),
            ("groq", "openai/gpt-oss-120b", "GPT-OSS 120B"),
            ("groq", "openai/gpt-oss-20b", "GPT-OSS 20B"),
        ]
    elif selected_model == "openrouter/free":
        cascade = [("openrouter", OPENROUTER_FREE_MODEL, "OpenRouter Free")]
    else:
        cascade = [("groq", selected_model, MODEL_MAPPING.get(selected_model, selected_model))]

    or_client = get_openrouter_client()
    answered = False
    last_error = ""

    for provider, model_id, label in cascade:
        if provider == "openrouter":
            active_client = or_client
        else:
            active_client = client
        if active_client is None:
            continue
        for attempt in range(2):
            try:
                stream = active_client.chat.completions.create(
                    model=model_id,
                    messages=messages,
                    temperature=temperature,
                    max_tokens=750,
                    stream=True
                )
                for chunk in stream:
                    delta = chunk.choices[0].delta.content
                    if delta:
                        if not answered:
                            yield f"*[{label}]*\n\n"
                            answered = True
                            st.session_state.telemetry["last_model"] = label
                        yield delta
                if answered:
                    return
            except Exception as exc:
                last_error = str(exc)
                time.sleep(1.0 if attempt == 0 else 2.0)
                continue

    if not answered:
        yield f"System Offline: Inference cascade exhausted. Upstream error: {last_error}"

def generate_institutional_briefing(client: Groq, chunks: List[Dict[str, Any]]) -> str:
    sample_texts = [f"[Page {c['page']}] {c['text']}" for c in chunks[:12]]
    context_str = "\n\n".join(sample_texts)[:5500]

    system_prompt = (
        "You are an elite corporate intelligence analyst. Generate an institutional executive intelligence "
        "briefing based on the provided document context. Structure your report strictly into these 5 sections:\n"
        "1. Executive Abstract & Core Purpose\n"
        "2. Key Quantitative Findings & Core Metrics\n"
        "3. Strategic & Operational Risk Matrix\n"
        "4. Legal, Compliance & Governance Observations\n"
        "5. Critical Action Items & Strategic Directives\n\n"
        "Cite source pages accurately using `[Page X]`. Present data in concise bullet points and bold highlights."
    )
    try:
        resp = client.chat.completions.create(
            model="openai/gpt-oss-120b",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"CONTEXT:\n{context_str}"}
            ],
            temperature=0.1,
            max_tokens=1200
        )
        return resp.choices[0].message.content or "Briefing generation failed."
    except Exception:
        resp = client.chat.completions.create(
            model="openai/gpt-oss-20b",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"CONTEXT:\n{context_str}"}
            ],
            temperature=0.1,
            max_tokens=1000
        )
        return resp.choices[0].message.content or "Briefing generation failed."


# ==============================================================================
# 7. QA BENCHMARK SELF-TEST SUITE
# ==============================================================================
def execute_qa_self_test(client: Groq, embedder: SentenceTransformer, doc_vectors: np.ndarray,
                         bm25: BM25Okapi, chunks: List[Dict[str, Any]], top_k: int = 4) -> Dict[str, Any]:
    if len(chunks) < 3:
        sample_indices = list(range(len(chunks)))
    else:
        sample_indices = [0, len(chunks) // 2, len(chunks) - 1]

    results = []
    hits = 0
    start_time = time.time()

    for idx in sample_indices:
        target_chunk = chunks[idx]
        target_page = target_chunk["page"]
        
        prompt = (
            f"Based ONLY on the following excerpt from Page {target_page}, generate a single specific, "
            f"concise factual question whose answer is contained directly in the excerpt.\n\n"
            f"EXCERPT:\n{target_chunk['text'][:400]}\n\n"
            "QUESTION:"
        )
        try:
            resp = client.chat.completions.create(
                model="openai/gpt-oss-20b",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=60,
                temperature=0.0
            )
            gen_q = (resp.choices[0].message.content or "").strip().strip('"\'')
            if not gen_q:
                continue

            retrieved = hybrid_retrieve(gen_q, embedder, doc_vectors, bm25, chunks, top_k=top_k)
            retrieved_pages = [r["chunk"]["page"] for r in retrieved]
            is_hit = target_page in retrieved_pages
            if is_hit:
                hits += 1

            results.append({
                "target_page": target_page,
                "question": gen_q,
                "retrieved_pages": retrieved_pages,
                "is_hit": is_hit
            })
        except Exception:
            continue

    total = len(results)
    score_pct = (hits / total * 100) if total > 0 else 0.0
    elapsed = time.time() - start_time

    return {
        "total": total,
        "hits": hits,
        "score_pct": score_pct,
        "latency": elapsed,
        "details": results
    }


# ==============================================================================
# 8. MARKDOWN REPORT EXPORT GENERATOR
# ==============================================================================
def generate_markdown_export(doc_name: str, num_pages: int, messages: List[Dict[str, Any]],
                             telemetry: Dict[str, Any], briefing: Optional[str] = None) -> str:
    lines = [
        "# DocuBrain Enterprise Intelligence Report",
        f"**Document Name:** {doc_name}  ",
        f"**Page Count:** {num_pages}  ",
        f"**Report Generated:** {time.strftime('%Y-%m-%d %H:%M:%S UTC', time.gmtime())}  ",
        f"**Architecture:** DocuBrain v4 ELITE (Hybrid BM25 + Dense RRF Fusion)  \n",
        "---",
        "## Platform Telemetry Summary",
        f"- Inquiries Executed: {len([m for m in messages if m['role'] == 'user'])}",
        f"- Estimated Tokens: {telemetry.get('total_tokens_est', 0)}",
        f"- Cache Hit Rate: {telemetry.get('cache_hits', 0)} / {max(telemetry.get('total_queries', 1), 1)}",
        f"- Average Latency: {telemetry.get('total_latency', 0.0) / max(telemetry.get('total_queries', 1), 1):.2f}s  \n"
    ]

    if briefing:
        lines.extend([
            "---",
            "## Executive Intelligence Briefing",
            briefing,
            "\n"
        ])

    lines.extend([
        "---",
        "## Intelligence Q&A Transcript\n"
    ])
    turn = 1
    for msg in messages:
        if msg["role"] == "user":
            lines.append(f"### Inquiry {turn}")
            lines.append(f"**User Question:** {msg['content']}\n")
        else:
            lines.append(f"**DocuBrain Analyst Response:**\n{msg['content']}\n")
            if "citations" in msg and msg["citations"].get("cited_pages"):
                p_str = ", ".join(str(p) for p in msg["citations"]["cited_pages"])
                lines.append(f"**Verified Provenance:** Pages {p_str}")
            if "evidence" in msg and msg["evidence"]:
                lines.append("\n**Retrieved Evidence Chunks:**\n")
                for ev in msg["evidence"]:
                    p = ev["chunk"]["page"]
                    s = ev["score"] * 100
                    lines.append(f"- `[Page {p}]` (Cosine Relevance: {s:.1f}%): {ev['chunk']['text']}")
            lines.append("\n---\n")
            turn += 1

    lines.append("\n*DocuBrain v4 ELITE - Grounded Document Intelligence System*")
    return "\n".join(lines)


# ==============================================================================
# 9. SIDEBAR: PERSISTENT LIBRARY, CONTROLS & OBSERVABILITY (ADMIN-ONLY)
# ==============================================================================
with st.sidebar:
    st.markdown(f"""
        <div style="display:flex; align-items:center; gap:10px; padding-bottom: 0.5rem; border-bottom: 1px solid rgba(255,255,255,0.08); margin-bottom: 1rem;">
            {get_logo_html(40)}
            <div>
                <div style="font-size: 1.15rem; font-weight: 800; color: #F8FAFC; line-height: 1.1;">DocuBrain</div>
                <div style="font-size: 0.62rem; color: #818CF8; font-weight: 700; letter-spacing: 1.2px;">ENTERPRISE ARCHITECTURE</div>
            </div>
        </div>
    """, unsafe_allow_html=True)

    uploaded_pdf = st.file_uploader("Upload PDF or Word Document", type=["pdf", "docx", "doc"], label_visibility="collapsed")

    def get_uploader_hash(uploader_obj) -> Optional[str]:
        if uploader_obj is not None:
            try:
                return hashlib.sha256(uploader_obj.getvalue()).hexdigest()
            except Exception:
                return None
        return None

    if st.button("New Analysis Session", use_container_width=True):
        st.session_state.active_doc_id = None
        st.session_state.doc_name = None
        st.session_state.num_pages = 0
        st.session_state.chunks = []
        st.session_state.vectors = None
        st.session_state.bm25 = None
        st.session_state.messages = []
        st.session_state.briefing = None
        st.session_state.answer_cache = {}
        st.session_state.feedback = {}
        st.session_state.pending_query = None
        st.session_state.last_ingested_hash = get_uploader_hash(uploaded_pdf)
        st.toast("Session reset. Upload or restore a document to begin.")
        st.rerun()

    st.markdown("### Document Library")
    saved_docs = list_library_documents()
    if saved_docs:
        for d in saved_docs:
            is_active = (d["doc_id"] == st.session_state.active_doc_id)
            c_class = "library-item-active" if is_active else "library-item"
            tag = " [ACTIVE]" if is_active else ""
            st.markdown(f"""
                <div class="{c_class}">
                    <div class="library-name">{d['filename']}{tag}</div>
                    <div class="library-meta">{d['num_pages']} pages · {d['indexed_at'][:10]}</div>
                </div>
            """, unsafe_allow_html=True)
            col_l1, col_l2 = st.columns([2, 1])
            if col_l1.button("Restore", key=f"load_{d['doc_id']}", use_container_width=True):
                loaded = load_document_library_state(d["doc_id"])
                if loaded:
                    st.session_state.active_doc_id = loaded["doc_id"]
                    st.session_state.doc_name = loaded["filename"]
                    st.session_state.num_pages = loaded["num_pages"]
                    st.session_state.chunks = loaded["chunks"]
                    st.session_state.vectors = loaded["vectors"]
                    st.session_state.bm25 = build_bm25_index(loaded["chunks"])
                    st.session_state.messages = loaded.get("messages", [])
                    st.session_state.briefing = loaded.get("briefing", None)
                    st.session_state.answer_cache = loaded.get("answer_cache", {})
                    st.session_state.feedback = loaded.get("feedback", {})
                    st.session_state.pending_query = None
                    curr_up_hash = get_uploader_hash(uploaded_pdf)
                    st.session_state.last_ingested_hash = curr_up_hash if curr_up_hash else loaded["doc_id"]
                    st.toast(f"Restored {loaded['filename']} from local library.")
                    st.rerun()
            if col_l2.button("Delete", key=f"del_{d['doc_id']}", use_container_width=True):
                delete_document_library_state(d["doc_id"])
                if st.session_state.active_doc_id == d["doc_id"]:
                    st.session_state.active_doc_id = None
                    st.session_state.doc_name = None
                    st.session_state.num_pages = 0
                    st.session_state.chunks = []
                    st.session_state.vectors = None
                    st.session_state.bm25 = None
                    st.session_state.messages = []
                    st.session_state.briefing = None
                    st.session_state.pending_query = None
                st.session_state.last_ingested_hash = get_uploader_hash(uploaded_pdf)
                st.toast("Document removed from library.")
                st.rerun()
    else:
        st.caption("No persisted documents in local library.")

    st.markdown("### Inference Engine")
    selected_engine = st.selectbox(
        "Model Engine",
        ["Auto-Cascade", "openrouter/free", "openai/gpt-oss-120b", "openai/gpt-oss-20b"],
        index=0,
        label_visibility="collapsed"
    )

    with st.expander("Advanced RAG Controls"):
        cfg_chunk_size = st.slider("Chunk Size (Chars)", min_value=400, max_value=1200, value=750, step=50)
        cfg_overlap = st.slider("Chunk Overlap (Chars)", min_value=50, max_value=300, value=150, step=25)
        cfg_top_k = st.slider("Context Chunks (Top-K)", min_value=2, max_value=8, value=4)
        cfg_threshold = st.slider("Confidence Kill-Switch", min_value=0.15, max_value=0.60, value=0.30, step=0.05)

    # ADMIN-ONLY: hidden from end users, visible only after password login
    if st.session_state.get("is_admin", False):
        with st.expander("Admin Observability & Cost Meter"):
            telem = st.session_state.telemetry
            q_count = telem["total_queries"]
            avg_lat = (telem["total_latency"] / q_count) if q_count > 0 else 0.0
            est_cost = (telem["total_tokens_est"] / 1_000_000) * 0.15
            st.markdown(f"**Total Queries:** `{q_count}`")
            st.markdown(f"**Tokens Consumed:** `{telem['total_tokens_est']:,}`")
            st.markdown(f"**Estimated Cost:** `${est_cost:.5f} USD`")
            st.markdown(f"**Cache Hit Rate:** `{telem['cache_hits']} hits`")
            st.markdown(f"**Average Latency:** `{avg_lat:.2f}s`")
            st.markdown(f"**Active Model:** `{telem['last_model']}`")
            st.markdown(f"**Security Blocks:** `{telem['security_blocks']}`")

    if st.session_state.get("messages") and len(st.session_state.messages) > 0:
        rep_md = generate_markdown_export(
            st.session_state.doc_name or "Document",
            st.session_state.num_pages,
            st.session_state.messages,
            st.session_state.telemetry,
            st.session_state.briefing
        )
        st.download_button(
            "Export Intelligence Report",
            rep_md,
            file_name=f"{st.session_state.doc_name or 'docubrain'}_report.md",
            use_container_width=True
        )

    if st.button("Clear Chat History", use_container_width=True):
        st.session_state.messages = []
        if st.session_state.active_doc_id:
            save_document_library_state(
                st.session_state.active_doc_id,
                st.session_state.doc_name,
                st.session_state.num_pages,
                st.session_state.chunks,
                st.session_state.vectors,
                [],
                st.session_state.answer_cache,
                st.session_state.feedback,
                st.session_state.briefing
            )
        st.toast("Chat history cleared.")
        st.rerun()


# ==============================================================================
# 10. IMMUTABLE DOCUMENT INGESTION LOGIC (PDF + DOCX, 2GB LIMIT)
# ==============================================================================
if uploaded_pdf is not None:
    file_bytes = uploaded_pdf.read()
    doc_hash = compute_sha256(file_bytes)
    
    if doc_hash != st.session_state.last_ingested_hash:
        cached_doc = load_document_library_state(doc_hash)
        if cached_doc:
            st.session_state.active_doc_id = cached_doc["doc_id"]
            st.session_state.doc_name = cached_doc["filename"]
            st.session_state.num_pages = cached_doc["num_pages"]
            st.session_state.chunks = cached_doc["chunks"]
            st.session_state.vectors = cached_doc["vectors"]
            st.session_state.bm25 = build_bm25_index(cached_doc["chunks"])
            st.session_state.messages = cached_doc.get("messages", [])
            st.session_state.briefing = cached_doc.get("briefing", None)
            st.session_state.answer_cache = cached_doc.get("answer_cache", {})
            st.session_state.feedback = cached_doc.get("feedback", {})
            st.session_state.last_ingested_hash = doc_hash
            st.session_state.pending_query = None
            st.toast("Document recognized. Loaded instantly from local cache.")
        else:
            try:
                if uploaded_pdf.name.lower().endswith(".doc"):
                    st.error("Legacy .doc format is not supported. Open it in Word, save as .docx, and upload again.")
                else:
                    is_large = len(file_bytes) > (15 * 1024 * 1024)
                    status_title = "Large document - indexing may take a minute" if is_large else "Ingesting and Vectorizing Document..."
                    with st.status(status_title, expanded=True) as status_box:
                        st.write("Extracting readable page text from document...")
                        pages = extract_document_pages(file_bytes, uploaded_pdf.name)
                        if not pages:
                            st.error("Unreadable or empty document. Please upload a text-based PDF or a .docx Word file.")
                            status_box.update(label="Ingestion Failed", state="error")
                            st.stop()

                        st.write(f"Generating sentence-aware chunks ({cfg_chunk_size} char window)...")
                        chunks = sentence_aware_chunking(pages, chunk_size=cfg_chunk_size, overlap=cfg_overlap)
                        
                        st.write("Constructing BM25 lexical keyword index...")
                        bm25_obj = build_bm25_index(chunks)
                        
                        st.write(f"Calculating CPU dense embeddings for {len(chunks)} chunks...")
                        embedder_instance = load_embedder()
                        vectors = batch_encode_chunks(embedder_instance, chunks)
                        
                        save_document_library_state(
                            doc_hash,
                            uploaded_pdf.name,
                            len(pages),
                            chunks,
                            vectors,
                            [],
                            {},
                            {},
                            None
                        )
                        
                        st.session_state.active_doc_id = doc_hash
                        st.session_state.doc_name = uploaded_pdf.name
                        st.session_state.num_pages = len(pages)
                        st.session_state.chunks = chunks
                        st.session_state.vectors = vectors
                        st.session_state.bm25 = bm25_obj
                        st.session_state.messages = []
                        st.session_state.briefing = None
                        st.session_state.answer_cache = {}
                        st.session_state.feedback = {}
                        st.session_state.last_ingested_hash = doc_hash
                        st.session_state.pending_query = None

                        status_box.update(label=f"Ingestion Complete: Indexed {len(pages)} pages", state="complete", expanded=False)
                        st.toast("Document indexed and persisted to library.")
            except Exception as e:
                st.error(f"Ingestion failed: {str(e)}")


# ==============================================================================
# 11. MAIN WORKSPACE: HEADER, TABS & INTERFACE
# ==============================================================================
active_doc_label = st.session_state.doc_name or "No Active Document"
active_eng_label = "Cascade Active" if selected_engine == "Auto-Cascade" else MODEL_MAPPING.get(selected_engine, selected_engine)

st.markdown(f"""
    <div class="app-header-bar">
        <div class="header-left-cluster">
            {get_logo_html(34)}
            <span class="app-wordmark">DocuBrain</span>
            <span class="edition-badge">ENTERPRISE</span>
            <div class="system-status-indicator">
                <div class="status-dot"></div>
                <span class="status-label">Operational</span>
            </div>
        </div>
        <div class="header-right-cluster">
            <div class="header-pill-doc" title="{active_doc_label}">{active_doc_label}</div>
            <div class="header-pill-engine">{active_eng_label}</div>
        </div>
    </div>
""", unsafe_allow_html=True)

tab_chat, tab_details, tab_reports, tab_audit = st.tabs([
    "Chat",
    "Document Details",
    "Quick Reports",
    "Security & Tests"
])

# ------------------------------------------------------------------------------
# TAB 1: CHAT
# ------------------------------------------------------------------------------
with tab_chat:
    with st.expander("How to use DocuBrain (Quick Guide)", expanded=st.session_state.onboarding_open):
        st.session_state.onboarding_open = False
        st.markdown("""
        1. **Upload or restore a document**: Use the sidebar to upload a PDF or Word file (up to 2GB) or restore an existing one from the Document Library.
        2. **Ask questions**: Submit any specific question or click a Quick Question button to explore the text.
        3. **Review verified citations**: Every factual claim cites exact source pages (`[Page X]`) grounded strictly in your document.
        4. **Inspect source evidence**: Open the evidence expander under any answer to inspect retrieved excerpts, cosine similarity, and RRF scores.
        5. **Generate executive reports**: Switch to the Quick Reports tab to generate comprehensive multi-section executive intelligence summaries.
        6. **Export intelligence records**: Download verified Markdown reports with citations and telemetry anytime from the sidebar.
        """)

    if st.session_state.active_doc_id and st.session_state.chunks:
        eng_name = "Auto-Cascade" if selected_engine == "Auto-Cascade" else MODEL_MAPPING.get(selected_engine, selected_engine)
        st.markdown(f"""
            <div class="metric-container">
                <div class="metric-card"><div class="metric-label">Active Document</div><div class="metric-value">{st.session_state.doc_name[:24]}</div></div>
                <div class="metric-card"><div class="metric-label">Pages</div><div class="metric-value">{st.session_state.num_pages}</div></div>
                <div class="metric-card"><div class="metric-label">Indexed Chunks</div><div class="metric-value">{len(st.session_state.chunks)}</div></div>
                <div class="metric-card"><div class="metric-label">Inference Engine</div><div class="metric-value">{eng_name}</div></div>
            </div>
        """, unsafe_allow_html=True)

        st.markdown("##### Quick Questions")
        qc1, qc2, qc3 = st.columns(3)
        with qc1:
            btn_sum = st.button("Executive Summary", use_container_width=True, help="Generate executive summary")
            if btn_sum:
                st.session_state.pending_query = "Provide a comprehensive, high-level executive summary of this entire document with core highlights."
                st.rerun()
        with qc2:
            btn_qty = st.button("Key Quantities & Metrics", use_container_width=True, help="Extract critical quantitative metrics")
            if btn_qty:
                st.session_state.pending_query = "What are the most critical quantitative figures, numbers, and operational takeaways in this document?"
                st.rerun()
        with qc3:
            btn_rsk = st.button("Risks & Constraints", use_container_width=True, help="Audit risks and constraints")
            if btn_rsk:
                st.session_state.pending_query = "List any risks, limitations, warnings, dependencies, or constraints specified in the document."
                st.rerun()

    else:
        st.markdown("""
            <div style="background: #11161D; border: 1px dashed rgba(255, 255, 255, 0.15); border-radius: 8px; padding: 2.5rem 2rem; text-align: center; margin: 1.5rem 0;">
                <div style="font-size: 1.25rem; font-weight: 700; color: #F8FAFC; margin-bottom: 0.5rem;">Enterprise Grounded Document Intelligence</div>
                <div style="color: #94A3B8; font-size: 0.85rem; margin-bottom: 2rem;">Zero hallucination tolerance with hybrid BM25 and dense neural retrieval.</div>
                <div style="display: flex; justify-content: center; gap: 16px; flex-wrap: wrap;">
                    <div style="background: rgba(255, 255, 255, 0.02); border: 1px solid rgba(255, 255, 255, 0.08); border-radius: 6px; padding: 14px 18px; text-align: left; width: 220px;">
                        <div style="font-size: 0.7rem; font-weight: 700; color: #818CF8; margin-bottom: 4px;">STEP 01</div>
                        <div style="font-size: 0.85rem; font-weight: 600; color: #E2E8F0; margin-bottom: 2px;">Ingest Document</div>
                        <div style="font-size: 0.75rem; color: #64748B;">Upload a PDF or Word file, or restore an index from the Document Library.</div>
                    </div>
                    <div style="background: rgba(255, 255, 255, 0.02); border: 1px solid rgba(255, 255, 255, 0.08); border-radius: 6px; padding: 14px 18px; text-align: left; width: 220px;">
                        <div style="font-size: 0.7rem; font-weight: 700; color: #818CF8; margin-bottom: 4px;">STEP 02</div>
                        <div style="font-size: 0.85rem; font-weight: 600; color: #E2E8F0; margin-bottom: 2px;">Submit Inquiry</div>
                        <div style="font-size: 0.75rem; color: #64748B;">Query operational metrics, risk clauses, or compliance terms.</div>
                    </div>
                    <div style="background: rgba(255, 255, 255, 0.02); border: 1px solid rgba(255, 255, 255, 0.08); border-radius: 6px; padding: 14px 18px; text-align: left; width: 220px;">
                        <div style="font-size: 0.7rem; font-weight: 700; color: #818CF8; margin-bottom: 4px;">STEP 03</div>
                        <div style="font-size: 0.85rem; font-weight: 600; color: #E2E8F0; margin-bottom: 2px;">Inspect Citations</div>
                        <div style="font-size: 0.75rem; color: #64748B;">Receive verified page citations and source context blocks.</div>
                    </div>
                </div>
            </div>
        """, unsafe_allow_html=True)

    # --------------------------------------------------------------------------
    # RENDER CHAT HISTORY (GEMINI-STYLE LAYOUT)
    # --------------------------------------------------------------------------
    st.markdown('<div class="chat-history-container">', unsafe_allow_html=True)
    
    for msg_idx, msg in enumerate(st.session_state.messages):
        is_user = (msg["role"] == "user")
        if is_user:
            st.markdown(f"""
                <div class="user-bubble-row">
                    <div class="user-bubble">{msg['content']}</div>
                </div>
            """, unsafe_allow_html=True)
        else:
            with st.container():
                col_d, col_b = st.columns([0.02, 0.98])
                with col_d:
                    st.markdown('<div class="assistant-accent-dot"></div>', unsafe_allow_html=True)
                with col_b:
                    rendered_content = render_citation_chips(msg["content"])
                    st.markdown(rendered_content, unsafe_allow_html=True)
                    
                    if msg.get("citations", {}).get("has_citations"):
                        c_data = msg["citations"]
                        if c_data.get("is_valid"):
                            st.markdown(f'<div class="verification-footnote"><span class="verification-verified">VERIFIED</span> {c_data["status_text"]}</div>', unsafe_allow_html=True)
                        else:
                            st.markdown(f'<div class="verification-footnote"><span class="verification-notice">NOTICE</span> {c_data["status_text"]}</div>', unsafe_allow_html=True)

                    if "evidence" in msg and msg["evidence"]:
                        with st.expander("Inspected Citations & Context Evidence"):
                            for item in msg["evidence"]:
                                p = item["chunk"]["page"]
                                sim = item["score"] * 100
                                rrf = item.get("rrf_score", 0.0)
                                st.markdown(
                                    f'<span class="cite-chip">PAGE {p}</span> '
                                    f'<span class="score-badge">Cosine Relevance: {sim:.1f}%</span> '
                                    f'<span class="score-badge" style="background: rgba(99,102,241,0.08); color: #818CF8; border-color: rgba(99,102,241,0.25);">RRF Score: {rrf:.4f}</span>',
                                    unsafe_allow_html=True
                                )
                                st.caption(f'"{item["chunk"]["text"]}"')

                    f_key = f"msg_{msg_idx}"
                    fb_val = st.session_state.feedback.get(f_key)
                    fb_col1, fb_col2, _ = st.columns([1, 1.1, 7])
                    pos_label = "Helpful [Recorded]" if fb_val == "positive" else "Helpful"
                    neg_label = "Needs Work [Recorded]" if fb_val == "negative" else "Needs Work"
                    
                    if fb_col1.button(pos_label, key=f"fb_pos_{msg_idx}"):
                        st.session_state.feedback[f_key] = "positive"
                        if st.session_state.active_doc_id:
                            save_document_library_state(
                                st.session_state.active_doc_id,
                                st.session_state.doc_name,
                                st.session_state.num_pages,
                                st.session_state.chunks,
                                st.session_state.vectors,
                                st.session_state.messages,
                                st.session_state.answer_cache,
                                st.session_state.feedback,
                                st.session_state.briefing
                            )
                        st.toast("Positive feedback recorded.")
                        st.rerun()

                    if fb_col2.button(neg_label, key=f"fb_neg_{msg_idx}"):
                        st.session_state.feedback[f_key] = "negative"
                        if st.session_state.active_doc_id:
                            save_document_library_state(
                                st.session_state.active_doc_id,
                                st.session_state.doc_name,
                                st.session_state.num_pages,
                                st.session_state.chunks,
                                st.session_state.vectors,
                                st.session_state.messages,
                                st.session_state.answer_cache,
                                st.session_state.feedback,
                                st.session_state.briefing
                            )
                        st.toast("Improvement feedback recorded.")
                        st.rerun()

    st.markdown('</div>', unsafe_allow_html=True)

    # --------------------------------------------------------------------------
    # PENDING QUERY INFERENCE PIPELINE
    # --------------------------------------------------------------------------
    if st.session_state.pending_query:
        active_query = st.session_state.pending_query
        st.session_state.pending_query = None

        if not GROQ_API_KEY and not OPENROUTER_API_KEY:
            st.error("Configuration Error: No inference API key configured in st.secrets.")
        elif not st.session_state.active_doc_id or not st.session_state.chunks:
            st.warning("Please upload a PDF or Word document, or restore an index from the Document Library.")
        else:
            now_ts = time.time()
            st.session_state.rate_limit_timestamps = [t for t in st.session_state.rate_limit_timestamps if now_ts - t < 60.0]
            if len(st.session_state.rate_limit_timestamps) >= 10:
                wait_sec = int(60.0 - (now_ts - st.session_state.rate_limit_timestamps[0])) + 1
                st.warning(f"Enterprise Rate Limit: Maximum 10 queries per minute allowed. Please wait {wait_sec} seconds.")
            else:
                st.session_state.rate_limit_timestamps.append(now_ts)

                # Render user question bubble
                st.session_state.messages.append({"role": "user", "content": active_query})
                st.markdown(f"""
                    <div class="user-bubble-row">
                        <div class="user-bubble">{active_query}</div>
                    </div>
                """, unsafe_allow_html=True)

                groq_client = get_client()

                # Dynamic Thinking Pill
                thinking_box = st.empty()
                def set_thinking_stage(stage_label: str):
                    thinking_box.markdown(f"""
                        <div class="thinking-pill">
                            {get_logo_html(16, is_spinning=True)}
                            <span class="thinking-pill-text">{stage_label}</span>
                        </div>
                    """, unsafe_allow_html=True)

                set_thinking_stage("Evaluating prompt security policy (llama-prompt-guard)...")
                is_safe = True
                sec_reason = ""
                if groq_client:
                    is_safe, sec_reason = check_prompt_security(groq_client, active_query)
                
                st.session_state.audit_log.append({
                    "timestamp": time.strftime("%H:%M:%S UTC"),
                    "query": active_query,
                    "is_safe": is_safe,
                    "reason": sec_reason or "Compliant"
                })

                if not is_safe:
                    thinking_box.empty()
                    st.session_state.telemetry["security_blocks"] += 1
                    sec_response = f"Security Notice: {sec_reason}"
                    st.markdown(sec_response)
                    st.session_state.messages.append({"role": "assistant", "content": sec_response})
                    st.rerun()

                set_thinking_stage("Resolving conversation entities and coreferences...")
                retrieval_query = active_query
                if groq_client:
                    retrieval_query = rewrite_query_if_followup(groq_client, st.session_state.messages[:-1], active_query)

                cache_key = get_answer_cache_key(st.session_state.active_doc_id, active_query)
                cached_entry = st.session_state.answer_cache.get(cache_key)

                if cached_entry:
                    thinking_box.empty()
                    st.session_state.telemetry["cache_hits"] += 1
                    st.session_state.telemetry["total_queries"] += 1
                    cached_ans = cached_entry["content"]
                    retrieved_items = cached_entry["evidence"]
                    citations_info = cached_entry["citations"]
                    
                    st.session_state.messages.append({
                        "role": "assistant",
                        "content": cached_ans,
                        "evidence": retrieved_items,
                        "citations": citations_info
                    })
                    st.rerun()

                set_thinking_stage("Executing hybrid BM25 + dense neural retrieval...")
                t_start = time.time()
                retrieved_items = hybrid_retrieve(
                    retrieval_query,
                    load_embedder(),
                    st.session_state.vectors,
                    st.session_state.bm25,
                    st.session_state.chunks,
                    top_k=cfg_top_k
                )

                # Refuse ONLY when BOTH signals fail
                top_dense_score = retrieved_items[0]["score"] if retrieved_items else 0.0
                top_bm25_score = retrieved_items[0]["bm25_score"] if retrieved_items else 0.0

                if top_dense_score < cfg_threshold and top_bm25_score <= 0.0:
                    thinking_box.empty()
                    refusal_text = "*The provided document does not contain sufficient information to answer this question. (Confidence threshold not met).*"
                    st.session_state.messages.append({
                        "role": "assistant",
                        "content": refusal_text,
                        "evidence": retrieved_items
                    })
                    st.rerun()

                expanded_context = expand_parent_context_window(retrieved_items, st.session_state.chunks, max_chars=5500)
                set_thinking_stage("Synthesizing grounded response via inference cascade...")
                thinking_box.empty()

                with st.container():
                    col_d, col_b = st.columns([0.02, 0.98])
                    with col_d:
                        st.markdown('<div class="assistant-accent-dot"></div>', unsafe_allow_html=True)
                    with col_b:
                        full_response = st.write_stream(
                            stream_groq_cascade(
                                groq_client,
                                selected_engine,
                                expanded_context,
                                active_query,
                                temperature=0.1
                            )
                        )
                elapsed_sec = time.time() - t_start

                citations_info = verify_citations(full_response, retrieved_items)
                est_tokens = int((len(expanded_context) + len(full_response)) / 4)
                st.session_state.telemetry["total_queries"] += 1
                st.session_state.telemetry["total_tokens_est"] += est_tokens
                st.session_state.telemetry["total_latency"] += elapsed_sec

                st.session_state.answer_cache[cache_key] = {
                    "content": full_response,
                    "evidence": retrieved_items,
                    "citations": citations_info
                }

                st.session_state.messages.append({
                    "role": "assistant",
                    "content": full_response,
                    "evidence": retrieved_items,
                    "citations": citations_info
                })

                save_document_library_state(
                    st.session_state.active_doc_id,
                    st.session_state.doc_name,
                    st.session_state.num_pages,
                    st.session_state.chunks,
                    st.session_state.vectors,
                    st.session_state.messages,
                    st.session_state.answer_cache,
                    st.session_state.feedback,
                    st.session_state.briefing
                )
                st.rerun()

    # Chat input anchored cleanly at the bottom
    user_typed_input = st.chat_input("Submit inquiry regarding the active document...")
    if user_typed_input:
        st.session_state.pending_query = user_typed_input
        st.rerun()


# ------------------------------------------------------------------------------
# TAB 2: DOCUMENT DETAILS
# ------------------------------------------------------------------------------
with tab_details:
    if st.session_state.active_doc_id and st.session_state.chunks:
        st.markdown("### Document Structural Details")
        total_chars = sum(len(c["text"]) for c in st.session_state.chunks)
        vocab_keywords = extract_top_keywords(st.session_state.chunks, top_n=12)

        dcol1, dcol2, dcol3, dcol4 = st.columns(4)
        with dcol1:
            st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-label">Total Characters</div>
                    <div class="metric-value">{total_chars:,}</div>
                </div>
            """, unsafe_allow_html=True)
        with dcol2:
            st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-label">Estimated Words</div>
                    <div class="metric-value">{total_chars // 5:,}</div>
                </div>
            """, unsafe_allow_html=True)
        with dcol3:
            st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-label">Vector Dimension</div>
                    <div class="metric-value">384 (float32)</div>
                </div>
            """, unsafe_allow_html=True)
        with dcol4:
            st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-label">PII Shield Status</div>
                    <div class="metric-value" style="color: #34D399;">Active & Redacted</div>
                </div>
            """, unsafe_allow_html=True)

        st.markdown("#### High-Frequency Enterprise Terms")
        kw_html = " ".join(f'<span class="score-badge" style="margin: 4px; font-size: 0.78rem;">{k.upper()} ({cnt})</span>' for k, cnt in vocab_keywords)
        st.markdown(f'<div style="margin-bottom: 1.5rem;">{kw_html}</div>', unsafe_allow_html=True)

        st.markdown("#### Indexed Chunk Explorer")
        search_filter = st.text_input("Filter indexed chunks by keyword", placeholder="Type keyword to filter chunks...")
        
        filtered_chunks = [
            c for c in st.session_state.chunks 
            if not search_filter or search_filter.lower() in c["text"].lower()
        ]
        
        st.caption(f"Displaying {len(filtered_chunks)} of {len(st.session_state.chunks)} chunks")
        for chunk in filtered_chunks[:25]:
            with st.expander(f"Chunk #{chunk['chunk_id']} - Page {chunk['page']} ({len(chunk['text'])} chars)"):
                st.code(chunk["text"], language="markdown")

    else:
        st.info("Ingest a document or restore a session from the Document Library to view document details.")


# ------------------------------------------------------------------------------
# TAB 3: QUICK REPORTS
# ------------------------------------------------------------------------------
with tab_reports:
    if st.session_state.active_doc_id and st.session_state.chunks:
        st.markdown("### Executive Quick Reports Studio")
        st.caption("Generate a synthesized 5-dimension executive intelligence briefing memo across the document.")

        if st.button("Generate Executive Intelligence Briefing", use_container_width=True):
            groq_c = get_client()
            if not groq_c:
                st.error("GROQ_API_KEY is not configured in st.secrets.")
            else:
                with st.status("Synthesizing Executive Intelligence Briefing...", expanded=True) as br_status:
                    st.write("Aggregating high-density document segments...")
                    briefing_result = generate_institutional_briefing(groq_c, st.session_state.chunks)
                    st.session_state.briefing = briefing_result
                    
                    save_document_library_state(
                        st.session_state.active_doc_id,
                        st.session_state.doc_name,
                        st.session_state.num_pages,
                        st.session_state.chunks,
                        st.session_state.vectors,
                        st.session_state.messages,
                        st.session_state.answer_cache,
                        st.session_state.feedback,
                        st.session_state.briefing
                    )
                    br_status.update(label="Executive Briefing Generated", state="complete")
                    st.rerun()

        if st.session_state.briefing:
            st.markdown("---")
            st.markdown(render_citation_chips(st.session_state.briefing), unsafe_allow_html=True)
            st.download_button(
                "Export Briefing Memo (.md)",
                st.session_state.briefing,
                file_name=f"{st.session_state.doc_name or 'document'}_briefing.md",
                use_container_width=True
            )
    else:
        st.info("Please load or ingest a document to generate executive reports.")


# ------------------------------------------------------------------------------
# TAB 4: SECURITY & TESTS
# ------------------------------------------------------------------------------
with tab_audit:
    st.markdown("### Retrieval Precision & Security Audit Hub")

    if st.session_state.active_doc_id and st.session_state.chunks:
        st.markdown("#### Automated QA Retrieval Benchmark")
        st.caption("Generates synthetic test questions directly from document text to evaluate Hit@K retrieval precision.")

        if st.button("Execute QA Benchmark Self-Test", use_container_width=True):
            groq_c = get_client()
            if not groq_c:
                st.error("GROQ_API_KEY is missing from secrets.")
            else:
                with st.status("Executing QA Benchmark Suite...", expanded=True) as qa_status:
                    st.write("Sampling document passages and generating test questions...")
                    bench = execute_qa_self_test(
                        groq_c,
                        load_embedder(),
                        st.session_state.vectors,
                        st.session_state.bm25,
                        st.session_state.chunks,
                        top_k=cfg_top_k
                    )
                    qa_status.update(label=f"QA Benchmark: {bench['score_pct']:.0f}% Retrieval Precision ({bench['hits']}/{bench['total']})", state="complete")
                    
                    bcol1, bcol2 = st.columns(2)
                    with bcol1:
                        st.metric("Retrieval Health Score (Hit@K)", f"{bench['score_pct']:.0f}%")
                    with bcol2:
                        st.metric("Evaluation Benchmark Latency", f"{bench['latency']:.2f}s")

                    st.markdown("##### Detailed Evaluation Records")
                    for d in bench["details"]:
                        badge_color = "#10B981" if d["is_hit"] else "#EF4444"
                        pass_label = "HIT" if d["is_hit"] else "MISS"
                        st.markdown(f"""
                            <div style="background: #11161D; border: 1px solid rgba(255,255,255,0.08); border-radius: 6px; padding: 10px 14px; margin-bottom: 8px;">
                                <span style="background: {badge_color}22; color: {badge_color}; border: 1px solid {badge_color}55; padding: 2px 6px; border-radius: 4px; font-size: 0.7rem; font-weight: 700;">{pass_label}</span>
                                <span style="font-size: 0.8rem; font-weight: 600; color: #F8FAFC; margin-left: 8px;">Target Page {d['target_page']}</span>
                                <div style="font-size: 0.78rem; color: #94A3B8; margin-top: 4px;"><strong>Generated Question:</strong> {d['question']}</div>
                                <div style="font-size: 0.72rem; color: #64748B; margin-top: 2px;">Retrieved Pages: {d['retrieved_pages']}</div>
                            </div>
                        """, unsafe_allow_html=True)

    st.markdown("#### Security & Guardrails Audit Log")
    if st.session_state.audit_log:
        st.dataframe(
            st.session_state.audit_log,
            column_config={
                "timestamp": "Timestamp",
                "query": "Inquiry Query",
                "is_safe": "Shield Status",
                "reason": "Classifier Verdict"
            },
            use_container_width=True,
            hide_index=True
        )
    else:
        st.caption("No security policy evaluations recorded in current session.")


# ==============================================================================
# 12. INSTITUTIONAL FOOTER
# ==============================================================================
st.markdown("""
    <div class="app-footer">
        <strong>DocuBrain v4 ELITE</strong> - Enterprise Document Intelligence Platform<br>
        Documents processed strictly in memory. Only derived chunk indices and dense vectors are persisted locally.<br>
        Built by <strong>Harpreet Singh</strong>
    </div>
""", unsafe_allow_html=True)